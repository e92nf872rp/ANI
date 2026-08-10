"""Verify heading detection for DOCX (DoclingReader) and PDF (PyMuPDF)."""
import os
import sys
import tempfile
from pathlib import Path

os.chdir(Path(__file__).resolve().parents[3])
sys.path.insert(0, ".")

from app.services.parse_service import ParseService, _emit_text_table_nodes

# ── 1. DOCX via DoclingReader ──
print("=" * 80)
print("[1] DOCX: Does DoclingReader emit ATX headings (# Title)?")
print("=" * 80)
from docx import Document

with tempfile.TemporaryDirectory() as tmpdir:
    docx_path = os.path.join(tmpdir, "heading_test.docx")
    doc = Document()
    doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("Intro paragraph.")
    doc.add_heading("Section A", level=2)
    doc.add_paragraph("Detail paragraph.")
    doc.save(docx_path)

    # Check raw DoclingReader output
    from llama_index.readers.docling import DoclingReader
    reader = DoclingReader()
    docs = reader.load_data(file_path=docx_path)
    raw_md = "\n\n".join(d.text for d in docs)
    print(f"Raw DoclingReader output:\n{repr(raw_md[:300])}")

    # Check parsed nodes
    svc = ParseService(uploader=None)
    nodes = svc.parse(docx_path)
    print(f"\nParsed nodes ({len(nodes)}):")
    for i, n in enumerate(nodes):
        sub = n.metadata.get("sub_type", "?")
        hl = n.metadata.get("heading_level", "-")
        sp = n.metadata.get("section_path", "")
        preview = n.content.replace("\n", "\\n")[:80]
        print(f"  Node[{i}] type={n.content_type} sub={sub} level={hl} path='{sp}' content: {preview}")

# ── 2. PDF via PyMuPDF ──
print("\n" + "=" * 80)
print("[2] PDF: Does PyMuPDF text contain heading markers?")
print("=" * 80)
import fitz

with tempfile.TemporaryDirectory() as tmpdir:
    pdf_path = os.path.join(tmpdir, "heading_test.pdf")
    doc = fitz.open()
    page = doc.new_page()
    # Insert text with different sizes to simulate headings
    page.insert_text((72, 72), "Chapter One", fontsize=20)  # Heading-like
    page.insert_text((72, 100), "Intro paragraph.", fontsize=12)
    page.insert_text((72, 130), "Section A", fontsize=16)  # Subheading-like
    page.insert_text((72, 160), "Detail paragraph.", fontsize=12)
    doc.save(pdf_path)
    doc.close()

    # Check raw PyMuPDF text
    doc = fitz.open(pdf_path)
    raw_text = doc[0].get_text("text")
    doc.close()
    print(f"Raw PyMuPDF text:\n{repr(raw_text)}")

    # Check parsed nodes
    svc = ParseService(uploader=None)
    nodes = svc.parse(pdf_path)
    print(f"\nParsed nodes ({len(nodes)}):")
    for i, n in enumerate(nodes):
        sub = n.metadata.get("sub_type", "?")
        hl = n.metadata.get("heading_level", "-")
        sp = n.metadata.get("section_path", "")
        preview = n.content.replace("\n", "\\n")[:80]
        print(f"  Node[{i}] type={n.content_type} sub={sub} level={hl} path='{sp}' content: {preview}")

print("\nDONE")
