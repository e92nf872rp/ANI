"""End-to-end functional test for parse_service (issue-009).

Tests against REAL MinIO (10.10.1.66:30900) and uses DoclingReader for
parsing. Creates test documents, parses them, and verifies:
  - AC2: DoclingReader parses MD/TXT with tables
  - AC4: Tables → HTML, large table splitting with header preservation
  - AC5: Image upload to MinIO, [图片: caption](OSS_URL) placeholder

Run from repo root: python -m pytest ai/rag-engine/tests/test_e2e_parse.py -v -s
"""
from __future__ import annotations

import os
from pathlib import Path

# Load .env so Settings picks up MINIO_ENDPOINT etc.
os.chdir(Path(__file__).resolve().parents[3])

import pytest
from app.clients.minio_client import ImageUploader
from app.services.parse_service import ParseService

# These are end-to-end integration tests that require a REAL MinIO server
# (10.10.1.66:30900) plus real docling/docx/PDF parsing libraries. They cannot
# run in the standard unit-test CI job, so they are skipped by default and
# executed explicitly with RUN_E2E=1 (e.g. a dedicated integration runner).
pytestmark = pytest.mark.skipif(
    os.environ.get("RUN_E2E") != "1",
    reason="e2e tests require real MinIO + parsing libraries (set RUN_E2E=1)",
)

# ── Fixtures ──────────────────────────────────────────────────────────────────


@pytest.fixture
def uploader():
    """Real MinIO ImageUploader (requires 10.10.1.66:30900 reachable)."""
    return ImageUploader()


@pytest.fixture
def service(uploader):
    """ParseService with real MinIO uploader."""
    return ParseService(uploader=uploader)


# ── AC2: DoclingReader parses TXT/MD ──────────────────────────────────────────


def test_e2e_parse_txt_with_table(service):
    """AC2: Parse a TXT file containing a markdown pipe table."""
    import tempfile

    content = "# Title\n\nSome intro text.\n\n| Name | Age |\n|---|---|\n| Alice | 30 |\n| Bob | 25 |\n\nMore text."
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
        f.write(content)
        f.flush()
        nodes = service.parse(f.name)

    assert len(nodes) > 0
    text_nodes = [n for n in nodes if n.content_type == "text"]
    table_nodes = [n for n in nodes if n.content_type == "table"]
    assert len(text_nodes) >= 1, "Should have text nodes"
    assert len(table_nodes) >= 1, "Should have at least one table node"
    assert "<table>" in table_nodes[0].content
    assert "<th>Name</th>" in table_nodes[0].content
    assert "<th>Age</th>" in table_nodes[0].content


def test_e2e_parse_md_with_table(service):
    """AC2: Parse a MD file with pipe table."""
    import tempfile

    content = "# Markdown Doc\n\n| Col1 | Col2 |\n|---|---|\n| a | b |\n| c | d |\n"
    with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
        f.write(content)
        f.flush()
        nodes = service.parse(f.name)

    table_nodes = [n for n in nodes if n.content_type == "table"]
    assert len(table_nodes) >= 1
    assert "<th>Col1</th>" in table_nodes[0].content
    assert "<th>Col2</th>" in table_nodes[0].content


# ── AC4: Large table splitting ────────────────────────────────────────────────


def test_e2e_large_table_split(service):
    """AC4: Table > 2048 tokens is split by row groups, each with header."""
    import tempfile

    # Build a large pipe table (each row ~200 chars, ~30 rows => >2048 tokens).
    header = "| ID | Description |"
    separator = "|---|---|"
    rows = [f"| {i} | {'x' * 190} |" for i in range(30)]
    content = "\n".join([header, separator] + rows)
    with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
        f.write(content)
        f.flush()
        nodes = service.parse(f.name)

    table_nodes = [n for n in nodes if n.content_type == "table"]
    assert len(table_nodes) > 1, f"Large table should be split into multiple nodes, got {len(table_nodes)}"
    # Each split table must contain the header.
    for tn in table_nodes:
        assert "<th>ID</th>" in tn.content, "Each split group must preserve header"
        assert "<th>Description</th>" in tn.content


# ── AC5: Image upload to MinIO ───────────────────────────────────────────────


def test_e2e_image_upload_to_minio(uploader):
    """AC5: ImageUploader uploads bytes to real MinIO and returns OSS URL."""
    test_data = b"\x89PNG fake image data for e2e test"
    oss_url = uploader.upload(test_data, "e2e-test/test-doc", "png")

    assert oss_url.startswith("ani-kb-docs/"), f"Expected bucket prefix, got: {oss_url}"
    assert "e2e-test/test-doc/images/" in oss_url
    assert oss_url.endswith(".png")


def test_e2e_image_placeholder_in_parse(service):
    """AC5: Embedded images in Word docs are uploaded and replaced with
    ``[图片: caption](OSS_URL)`` placeholder nodes.
    """
    import tempfile

    from docx import Document
    from docx.shared import Inches

    # Minimal valid 1x1 PNG (python-docx validates image headers).
    PNG_DATA = (
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
        b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xcf\xc0"
        b"\x00\x00\x00\x03\x00\x01\x5c\x9d\x87\x9b\x00\x00\x00\x00IEND\xaeB`\x82"
    )

    with tempfile.TemporaryDirectory() as tmpdir:
        # Create a real .docx with an embedded image.
        img_path = Path(tmpdir) / "test_img.png"
        img_path.write_bytes(PNG_DATA)

        docx_path = Path(tmpdir) / "img_test.docx"
        doc = Document()
        doc.add_paragraph("Text before image.")
        doc.add_picture(str(img_path), width=Inches(1))
        doc.add_paragraph("Text after image.")
        doc.save(str(docx_path))

        nodes = service.parse(str(docx_path), object_prefix="e2e-test/img")

        # Expect at least one image node with a MinIO OSS URL.
        image_nodes = [n for n in nodes if n.content_type == "image"]
        assert len(image_nodes) >= 1, f"Expected image node, got types: {[n.content_type for n in nodes]}"
        assert "[图片:" in image_nodes[0].content
        assert "ani-kb-docs/e2e-test/img/images/" in image_nodes[0].content


# ── Unsupported type ─────────────────────────────────────────────────────────


def test_e2e_unsupported_type_raises(service):
    """ParseService rejects unsupported file types."""
    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".html", delete=False) as f:
        f.write(b"<html></html>")
        f.flush()
        with pytest.raises(ValueError, match="doc.unsupported_type"):
            service.parse(f.name)


# ── PDF heading detection ────────────────────────────────────────────────────


def test_e2e_pdf_heading_detection(service):
    """PDF: large-font lines are detected as headings with section_path."""
    import tempfile

    import fitz

    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = str(Path(tmpdir) / "heading_test.pdf")
        doc = fitz.open()
        page = doc.new_page()
        # body=12, heading=20 (>=2x → ##), subheading=16 (>=1.5x → ###)
        page.insert_text((72, 72), "Chapter One", fontsize=20)
        page.insert_text((72, 100), "Intro paragraph.", fontsize=12)
        page.insert_text((72, 130), "Section A", fontsize=16)
        page.insert_text((72, 160), "Detail paragraph.", fontsize=12)
        doc.save(pdf_path)
        doc.close()

        nodes = service.parse(pdf_path)

        # Expect heading + paragraph nodes with section_path breadcrumbs.
        headings = [n for n in nodes if n.metadata.get("sub_type") == "heading"]
        assert len(headings) >= 2, f"Expected 2 headings, got {len(headings)}"

        # First heading: "Chapter One"
        assert "Chapter One" in headings[0].content
        assert headings[0].metadata["heading_level"] == 2  # 20 >= 12 * 1.5

        # Second heading: "Section A"
        assert "Section A" in headings[1].content
        assert headings[1].metadata["heading_level"] == 3  # 16 >= 12 * 1.25

        # Paragraph under Section A should carry the full section_path.
        section_text = next(n for n in nodes if "Detail paragraph" in n.content)
        assert section_text.metadata["section_path"] == "Chapter One > Section A"


# ── DOCX heading detection ──────────────────────────────────────────────────


def test_e2e_docx_heading_detection():
    """DOCX: Word headings are detected with section_path breadcrumbs.

    Uses _emit_text_table_nodes directly with the markdown that DoclingReader
    produces for Word headings, to avoid DoclingReader's internal document
    cache interfering with test isolation.
    """
    from app.services.parse_service import _emit_text_table_nodes

    # DoclingReader maps Word heading level 1 → "## ", level 2 → "### ".
    md = "## Chapter One\n\nIntro paragraph.\n\n### Section A\n\nDetail paragraph."
    nodes = _emit_text_table_nodes(md)

    headings = [n for n in nodes if n.metadata.get("sub_type") == "heading"]
    assert len(headings) >= 2, f"Expected 2+ headings, got {len(headings)}: {[(n.content, n.metadata) for n in nodes]}"

    assert "Chapter One" in headings[0].content
    assert headings[0].metadata["heading_level"] == 2

    assert "Section A" in headings[1].content
    assert headings[1].metadata["heading_level"] == 3

    # Paragraph under Section A carries section_path breadcrumb.
    detail = next(n for n in nodes if "Detail paragraph" in n.content)
    assert detail.metadata["section_path"] == "Chapter One > Section A"


if __name__ == "__main__":
    raise SystemExit(pytest.main([__file__, "-v", "-s"]))
