"""Debug: Check DOCX heading parsing through full service path."""
import os
import sys
import tempfile
from pathlib import Path

os.chdir(Path(__file__).resolve().parents[3])
sys.path.insert(0, ".")

from docx import Document
from app.services.parse_service import ParseService, _split_tables_and_text, _split_by_headings, _detect_heading

with tempfile.TemporaryDirectory() as tmpdir:
    docx_path = str(Path(tmpdir) / "heading_test.docx")
    doc = Document()
    doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("Intro paragraph.")
    doc.add_heading("Section A", level=2)
    doc.add_paragraph("Detail paragraph.")
    doc.save(docx_path)

    # Check raw DoclingReader output
    from llama_index.readers.docling import DoclingReader
    reader = DoclingReader()
    docs = reader.load_data(file_path=docx_path)
    raw_md = "\n\n".join(d.text for d in docs)
    print(f"Raw DoclingReader output:\n{repr(raw_md)}")
    print()

    # Check _split_tables_and_text
    segments = list(_split_tables_and_text(raw_md))
    print(f"_split_tables_and_text segments ({len(segments)}):")
    for i, (kind, seg) in enumerate(segments):
        print(f"  [{i}] kind={kind} seg={repr(seg[:100])}")
    print()

    # Check _split_by_headings on first text segment
    for kind, seg in segments:
        if kind == "text":
            sub_segs = _split_by_headings(seg)
            print(f"_split_by_headings on text segment ({len(sub_segs)}):")
            for i, s in enumerate(sub_segs):
                print(f"  [{i}] {repr(s[:80])}")
                h = _detect_heading(s)
                print(f"       heading={h}")
    print()

    # Full parse
    svc = ParseService(uploader=None)
    nodes = svc.parse(docx_path)
    print(f"Parsed nodes ({len(nodes)}):")
    for i, n in enumerate(nodes):
        sub = n.metadata.get("sub_type", "?")
        hl = n.metadata.get("heading_level", "-")
        sp = n.metadata.get("section_path", "")
        print(f"  Node[{i}] type={n.content_type} sub={sub} level={hl} path='{sp}' content: {repr(n.content[:80])}")

print("\nDONE")
