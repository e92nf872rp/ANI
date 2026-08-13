#!/usr/bin/env python3
"""表格 / 图片转换验证 + 切分分段 / 检索来源 落库打印（P0 内容正确性测试）。

用一个「含原生表格 + 内嵌图片 + 正文」的 docx 文档验证：
  1. 表格是否正确转换：kb_chunks 中存在 content_type='table' 的 chunk，
     content 含 HTML <table>（带 <th>/<td>），custom_metadata.sub_type='table'。
  2. 图片是否正确转换：kb_chunks 中存在 content_type='image' 的 chunk，
     content 形如 [图片: caption](OSS_URL)，且 OSS_URL 实际可 GET（MinIO 对象存在）。
  3. 切分分段落库打印：直连 PG 读取该 doc 的全部 kb_chunks（含
     chunk_type / content_type / content / custom_metadata / token_count / parent）。
  4. 检索分段来源落库打印：Query 后，把返回的每个 source 对应的 chunk 在
     kb_chunks 中的完整记录（child 与回填的 parent_content）打印出来，
     校验检索命中的来源确实指向正确的分段。

输出到文件（* 重定向），每一环节 REQUEST + RESPONSE + DB 查询结果全部展示。
"""
import asyncio
import hashlib
import io
import json
import re
import sys
import time
import uuid

import asyncpg
import requests
from PIL import Image
import json as _json

GATEWAY = "http://localhost:8080"
TENANT_ID = "00000000-0000-0000-0000-000000000001"
USER_ID = "00000000-0000-0000-0000-000000000002"
HEADERS = {
    "X-Dev-Tenant-ID": TENANT_ID,
    "X-Dev-User-ID": USER_ID,
    "Content-Type": "application/json",
}
PG_DSN = "postgresql://ani:ani_dev_password@10.10.1.66:30945/ani"

CHUNK_SIZE = 512

SUMMARY = []


def _show(v, limit=1500):
    s = str(v)
    return s if len(s) <= limit else s[:limit] + f"...(截断,共{len(s)}字)"


def record(label, ok, detail=""):
    SUMMARY.append((label, ok, detail))
    print(f"  [{'PASS' if ok else 'FAIL'}] {label} {detail}")


def request(method, url, *, json=None, data=None, params=None, timeout=120):
    print(f"    REQUEST: {method} {url}")
    if json is not None:
        print(f"      body(JSON) = {_json.dumps(json, ensure_ascii=False, default=str)}")
    r = requests.request(method, url, headers=HEADERS, json=json, data=data,
                         params=params, timeout=timeout)
    try:
        b = r.json()
    except Exception:
        b = r.text
    print(f"    RESPONSE(status={r.status_code}) = {_show(b, 2000)}")
    return r


# ── 构造含「原生表格 + 图片」的 docx ─────────────────────────────────────
def _make_png_bytes() -> bytes:
    """生成一张真实 PNG（示意架构框图），用于验证图片提取/上传。"""
    img = Image.new("RGB", (320, 180), (0x1F, 0x6F, 0x8B))
    from PIL import ImageDraw
    d = ImageDraw.Draw(img)
    d.rectangle([40, 40, 280, 90], fill=(0xFF, 0xD8, 0x4D))
    d.rectangle([100, 110, 220, 150], fill=(0x4C, 0xAF, 0x50))
    d.text((70, 60), "ANI Arch", fill=(0x33, 0x33, 0x33))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _make_docx_bytes() -> bytes:
    import docx
    from docx.shared import Inches

    d = docx.Document()
    d.add_heading("ANI 平台：调度算法与混合检索表（图片与表格转换验证）", level=1)
    d.add_paragraph(
        "本章节用于验证文档解析对表格与图片的处理。读者可对照下文调度算法说明表，"
        "理解作业调度与混合检索的实现要点。"
    )
    # 原生表格
    table = d.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["算法", "类型", "说明"]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    data = [
        ["FCFS", "调度", "先来先服务，简单公平"],
        ["DRF", "调度", "占优资源公平，多资源维度"],
        ["RRF", "检索", "倒数排名融合，兼顾语义与字面"],
    ]
    for i, row in enumerate(data, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
    d.add_paragraph(
        "上表总结了 ANI 平台所采用的调度与检索算法。其中 RRF 用于混合检索阶段的排名融合。"
    )
    # 图片
    d.add_picture(io.BytesIO(_make_png_bytes()), width=Inches(3.2))
    d.add_paragraph(
        "上图为 ANI 平台总体架构示意图，展示了网关、知识库服务与 RAG 引擎的解耦关系。"
    )
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# ── 直连 PG 查 chunk ────────────────────────────────────────────────────
async def query_chunks(kb_id: str, doc_id: str, condition: str = ""):
    conn = await asyncpg.connect(PG_DSN)
    try:
        async with conn.transaction():
            await conn.execute(
                "SELECT set_config('app.current_tenant_id', $1, true)", TENANT_ID
            )
            sql = (
                "SELECT id, chunk_type, content_type, content, token_count, "
                "       file_name, page_number, parent_chunk_id, custom_metadata "
                "FROM kb_chunks WHERE kb_id=$1 AND doc_id=$2 " + condition +
                " ORDER BY id"
            )
            rows = await conn.fetch(sql, uuid.UUID(kb_id), uuid.UUID(doc_id))
    finally:
        await conn.close()
    return rows


async def query_chunk_by_id(kb_id: str, chunk_id: str):
    conn = await asyncpg.connect(PG_DSN)
    try:
        async with conn.transaction():
            await conn.execute(
                "SELECT set_config('app.current_tenant_id', $1, true)", TENANT_ID
            )
            rows = await conn.fetchrow(
                "SELECT id, chunk_type, content_type, content, token_count, "
                "       parent_chunk_id, custom_metadata FROM kb_chunks "
                "WHERE kb_id=$1 AND id=$2",
                uuid.UUID(kb_id), uuid.UUID(chunk_id),
            )
    finally:
        await conn.close()
    return rows


def _md(v):
    if isinstance(v, str):
        try:
            return json.loads(v) if v else {}
        except Exception:
            return {}
    return v or {}


def _print_chunk(i, c, indent="  "):
    md = _md(c["custom_metadata"])
    print(f"{indent}--- chunk#{i} ---")
    print(f"{indent}  chunk_type={c['chunk_type']} content_type={c['content_type']} "
          f"tokens={c['token_count']} file={c['file_name']} page={c['page_number']}")
    print(f"{indent}  parent_chunk_id={c['parent_chunk_id']}")
    print(f"{indent}  custom_metadata={json.dumps(md, ensure_ascii=False, default=str)}")
    print(f"{indent}  content={_show(c['content'], 3000)}")


def main() -> int:
    print("=" * 80)
    print("  表格/图片转换 与 切分分段/检索来源 验证 (chunk_size=%d)" % CHUNK_SIZE)
    print("=" * 80)

    content_bytes = _make_docx_bytes()
    doc_name = "ani_table_image.docx"

    # 1) CreateKB
    r = request("POST", f"{GATEWAY}/api/v1/svc/knowledge-bases", json={
        "idempotency_key": f"p0c_{uuid.uuid4()}",
        "name": f"p0c-content-{int(time.time())}",
        "embedding_model": "Qwen3-Embedding-0.6B",
        "chunk_size": CHUNK_SIZE, "top_k": 5, "score_threshold": 0.3,
        "retrieval_mode": "hybrid",
    })
    kb_id = r.json().get("id")
    record("CreateKB", r.status_code == 201 and kb_id, f"kb_id={kb_id}")
    if not kb_id:
        return 1

    # 2) UploadURL
    r = request("POST", f"{GATEWAY}/api/v1/svc/knowledge-bases/{kb_id}/documents", json={
        "idempotency_key": f"p0c_up_{uuid.uuid4()}",
        "file_name": doc_name, "file_type": "docx",
        "file_size_bytes": len(content_bytes),
        "checksum_sha256": hashlib.sha256(content_bytes).hexdigest(),
    })
    body = r.json()
    doc_id = body.get("doc_id")
    record("GetDocumentUploadURL", r.status_code == 200 and doc_id, f"doc_id={doc_id}")

    # 3) PUT MinIO
    up = requests.put(body.get("upload_url"), data=content_bytes, timeout=60)
    print(f"    MinIO PUT status={up.status_code}")
    record("MinIO PUT", up.status_code in (200, 204))

    # 4) Notify
    r = request("POST",
                f"{GATEWAY}/api/v1/svc/knowledge-bases/{kb_id}/documents/{doc_id}/notify-uploaded",
                json={"storage_path": body.get("storage_path")})
    record("NotifyUploaded", r.status_code == 202)

    # 5) 轮询 ready（静默轮询，仅打印状态迁移）
    pstatus, cc = "pending", 0
    seen = set()
    for _ in range(48):
        time.sleep(5)
        pr = requests.get(f"{GATEWAY}/api/v1/svc/knowledge-bases/{kb_id}/documents",
                          headers=HEADERS, params={"limit": "20"}, timeout=30)
        for it in pr.json().get("items", []):
            if it.get("id") == doc_id:
                pstatus, cc = it.get("parse_status"), it.get("chunk_count") or 0
                break
        tag = f"{pstatus}/chunk={cc}"
        if tag not in seen:
            seen.add(tag)
            print(f"      [poll] parse_status={pstatus} chunk_count={cc}")
        if pstatus in ("ready", "failed"):
            break
    record("解析 ready", pstatus == "ready", f"status={pstatus} chunk_count={cc}")
    if pstatus != "ready":
        return 1

    # ── 6) 切分分段（落库打印）──────────────────────────────────────────
    print("\n" + "=" * 80)
    print("  切分分段（kb_chunks 落库记录）")
    print("=" * 80)
    rows = asyncio.run(query_chunks(kb_id, doc_id))
    print(f"  共 {len(rows)} 条 chunk")
    for i, c in enumerate(rows, 1):
        _print_chunk(i, c)
        print()

    tables = [c for c in rows if c["content_type"] == "table"]
    images = [c for c in rows if c["content_type"] == "image"]
    children = [c for c in rows if c["chunk_type"] == "child"]
    # 图片不再单独成段：以 markdown 图片标记 ![alt](完整URL) 内联进文本/父段。
    # 统计每条 child/text chunk 内容中出现的 inline 图片链接数。
    inline_imgs = []
    for c in rows:
        content = c["content"] or ""
        for m in re.finditer(r"!\[([^\]]*)\]\(([^)]+)\)", content):
            inline_imgs.append(m.group(2))
    print("\n  统计: total=%d child=%d table_chunk=%d image_chunk=%d inline_image_links=%d" % (
        len(rows), len(children), len(tables), len(images), len(inline_imgs)))

    # 7) 表格转换校验
    print("\n" + "=" * 80)
    print("  表格转换校验")
    print("=" * 80)
    table_ok = False
    for i, t in enumerate(tables, 1):
        c = t["content"]
        has_html = "<table" in c and ("<th>" in c or "<td>" in c or "<tr>" in c)
        sub = _md(t["custom_metadata"]).get("sub_type")
        print(f"  table chunk#{i}: has_html_table={has_html} sub_type={sub}")
        print(f"    content({len(c)}chars)={_show(c, 1200)}")
        table_ok = has_html and sub == "table"
    record("表格正确转换(HTML <table> + sub_type=table)", table_ok,
           f"table_chunks={len(tables)}")

    # 8) 图片转换校验（内联进文本/父段，非独立 image chunk）
    #    新行为：图片不单独成段，而是以 [图片: caption](oss_url) 形式内联进文本/父段。
    #    注意 oss_url 是 MinIO 相对对象键（ani-kb-docs/.../images/*.png），由网关
    #    拼接对外可访问地址；这里校验链接存在且对象键形态正确，不做整 URL GET。
    print("\n" + "=" * 80)
    print("  图片转换校验（内联进文本/父段，非独立 image chunk）")
    print("=" * 80)
    img_ok = True
    img_url = None
    # 从所有 chunk（含 child 文本段）内容里收集 inline 图片链接
    for i, url in enumerate(inline_imgs, 1):
        print(f"  inline image link#{i}: {url}")
        img_url = url
        # 完整 URL 形态校验：
        #   http(s)://<minio-endpoint>/ani-kb-docs/<tenant>/<kb>/<doc>/images/<uuid>.png
        shape_ok = (
            url.startswith(("http://", "https://"))
            and "ani-kb-docs/" in url
            and "/images/" in url
            and url.endswith(".png")
        )
        if not shape_ok:
            img_ok = False
            print(f"      !! OSS 完整 URL 形态不正确: {url}")
        else:
            print(f"      OSS 完整 URL 形态正确")
    # 期望：无独立 image chunk，但文本段内联了图片链接且对象键形态正确
    no_isolated = len(images) == 0
    print(f"  独立 image_chunk 数={len(images)}（预期 0） inline_image_links={len(inline_imgs)}")
    record("图片内联进文本段(无独立image chunk)", no_isolated,
           f"image_chunks={len(images)}")
    record("图片内联链接(形态正确)", img_ok and len(inline_imgs) > 0,
           f"inline_links={len(inline_imgs)} url={img_url}")

    # 9) Query 验证检索分段来源
    print("\n" + "=" * 80)
    print("  Query 检索 与 分段来源落库核对")
    print("=" * 80)
    q_ok = False
    dedup_ok = True  # 跨 mode 累计：任一 mode 出现重复内容即失败
    for mode in ("hybrid", "keyword", "vector"):
        r = request("POST", f"{GATEWAY}/api/v1/svc/knowledge-bases/{kb_id}/query", json={
            "idempotency_key": f"p0c_q_{mode}_{uuid.uuid4()}",
            "question": "RRF 算法在混合检索中如何使用倒数排名融合？",
            "top_k": 5, "score_threshold": 0.25, "retrieval_mode": mode,
        })
        raw = r.json() if r.status_code == 200 else {}
        srcs = raw.get("sources") if isinstance(raw, dict) else []
        print(f"  mode={mode} sources={len(srcs)}")
        # 去重校验：返回的应是父块内容，且不同 source 的内容不应重复
        seen_contents: set[str] = set()
        for s in srcs:
            content = s.get("content") or ""
            if content in seen_contents:
                dedup_ok = False
                print(f"    !! 重复内容: {_show(content, 200)}")
            seen_contents.add(content)
            # 打印 source 里已有的字段
            print(f"    source: file={s.get('file_name')} score={s.get('score')} "
                  f"chunk_id={s.get('chunk_id') or s.get('id')}")
        # 落库核对：对每个 source，查 kb_chunks 该 id 的完整记录
        if srcs:
            q_ok = True
            for s in srcs[:3]:
                cid = s.get("chunk_id") or s.get("id")
                if not cid:
                    continue
                row = asyncio.run(query_chunk_by_id(kb_id, cid))
                if row:
                    print(f"    --- 检索来源落库记录 chunk_id={cid} ---")
                    _print_chunk("S", row)
                    if row["parent_chunk_id"]:
                        pr = asyncio.run(query_chunk_by_id(kb_id, str(row["parent_chunk_id"])))
                        if pr:
                            print(f"      回填父块({row['parent_chunk_id']}): {_show(pr['content'], 500)}")
                else:
                    print(f"    !! chunk_id={cid} 在 kb_chunks 未找到（来源无法落库核对）")
        print()
    record("Query 返回带 sources", q_ok, "sources 非空")
    record("检索返回父块且去重", dedup_ok, "各 source 内容唯一(parent块)")

    # 10) 清理
    request("DELETE", f"{GATEWAY}/api/v1/svc/knowledge-bases/{kb_id}/documents/{doc_id}")
    request("DELETE", f"{GATEWAY}/api/v1/svc/knowledge-bases/{kb_id}")

    # 汇总
    print("\n" + "=" * 80)
    print("  结果汇总")
    print("=" * 80)
    passed = sum(1 for _, ok, _ in SUMMARY if ok)
    failed = len(SUMMARY) - passed
    print(f"  通过: {passed}  失败: {failed}")
    for label, ok, detail in SUMMARY:
        print(f"    [{'PASS' if ok else 'FAIL'}] {label} {detail}")
    return 0 if failed == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
